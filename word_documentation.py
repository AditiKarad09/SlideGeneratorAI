import json
import re
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import plot_code
from gemini_config import model
import PIL.Image

def generate_word_documentation(data):
    """
    1) Generate charts (plot_code.main()).
    2) Extract column names if data is a table.
    3) Build a Gemini prompt that asks for:
       - "Doc Title: <some name>"
       - multiple "Section X: Title" blocks
       - a full Markdown table if there's a table
       - references to any charts
    4) Parse Gemini's output to capture doc title, sections, bullet points,
       table lines, and chart references.
    5) Build a .docx with the doc title, headings, bullet points, table, charts.
    6) Save as output_document.docx.
    """

    # ----------------------------------------------------------------
    # 1) Generate or retrieve chart images
    # ----------------------------------------------------------------
    print("Generating charts for Word doc...")
    chart_file_list = plot_code.main()  # e.g. ["sales_plot.png", "temperature_plot.png", ...]
    data["chart_filename"] = chart_file_list

    # 2) Extract column names if 'type' is table
    columns_list = []
    if data.get("type") == "table" and "content" in data and data["content"]:
        # The first row's keys become our column names
        first_row = data["content"][0]
        columns_list = list(first_row.keys())  # e.g. ["Month", "Temperature (°C)", ...]
    else:
        columns_list = []

    # Build a string for columns
    # e.g. "Month, Temperature (°C), Rainfall (mm), Sales (Units), Advertising Budget ($)"
    col_str = ", ".join(columns_list) if columns_list else "No columns found"

    # If there's a document summary:
    doc_summary = data.get("document_summary", "")

    # ----------------------------------------------------------------
    # 3) Build the Gemini prompt
    # ----------------------------------------------------------------
    gemini_prompt = f"""
You are to produce a DETAILED Word document about the data and analysis below.

**Please output a top line**: 
Doc Title: <some relevant, descriptive title>

Then produce multiple sections using:
Section X: <Title>

**If the data is a table**, show the FULL table (all rows) in Markdown format:
- First row = column names
- Second row = dashes
- Then each row of data

if there are references present then list all the references do not say "copy from here" 
do not give "**" anywhere in the text 

When referencing charts, do:
Chart: <description> (filename.png)

Include bullet points if helpful.

**Data**:
{json.dumps(data, indent=4)}

**Extracted Columns**: {col_str}

**Document Summary**:
{doc_summary}

Example structure:
Doc Title: My Seasonal Sales Analysis
Section 1: Introduction
Content (no bullet points): ...
Bullet points:
- ...
| Month | Temperature | ...
|-------|------------| ...
| Jan   | 5          | ...
Chart: Explanation (temperature_plot.png)
...
"""

    try:
        response = model.generate_content(gemini_prompt)
        gemini_text = response.text.strip()
        print("[Gemini Output]\n", gemini_text)
    except Exception as e:
        print(f"Error calling Gemini: {str(e)}")
        return

    # ----------------------------------------------------------------
    # 4) Parse out "Doc Title:" from the top, then parse "Section X: Title"
    # ----------------------------------------------------------------
    doc_title_pattern = r"^Doc\s+Title\s*:\s*(.*)"
    title_match = re.search(doc_title_pattern, gemini_text, re.MULTILINE)
    if title_match:
        doc_title = title_match.group(1).strip()
        # Remove that line from the text
        gemini_text = re.sub(doc_title_pattern, "", gemini_text, count=1)
    else:
        doc_title = "Gemini Analysis Document"

    # Parse "Section X: Title"
    section_pattern = r"(Section\s+(\d+)\s*:\s*(.*?))(?=Section\s+\d+|$)"
    raw_blocks = re.findall(section_pattern, gemini_text, flags=re.DOTALL)

    # We'll store each section
    sections_data = []
    for full_block, section_num_str, after_colon_title in raw_blocks:
        try:
            section_num = int(section_num_str.strip())
        except ValueError:
            section_num = 999

        # lines[1:] => everything after "Section X: Title"
        lines = full_block.splitlines()
        block_lines = lines[1:]
        section_title = after_colon_title.split("\n")[0].strip()

        content_str = ""
        bullet_points = []
        chart_lines = []
        table_lines = []
        in_bullet_mode = False

        for line in block_lines:
            txt = line.strip()
            if not txt:
                in_bullet_mode = False
                continue

            content_match = re.match(r"^Content\s*\(no\s*bullet\s*points\)\s*:\s*(.*)$", txt, re.IGNORECASE)
            bulletpts_match = re.match(r"^Bullet\s*points\s*:\s*$", txt, re.IGNORECASE)
            chart_match = re.match(r"^-?\s*Chart:\s*(.*)$", txt, re.IGNORECASE)
            table_line_match = re.match(r"^\|.*\|.*\|", txt)  # naive check for markdown table

            if content_match:
                content_str += content_match.group(1).strip() + " "
                in_bullet_mode = False
            elif bulletpts_match:
                in_bullet_mode = True
            elif chart_match:
                chart_lines.append(chart_match.group(1).strip())
                in_bullet_mode = False
            elif table_line_match:
                table_lines.append(txt)
                in_bullet_mode = False
            else:
                if in_bullet_mode and txt.startswith("-"):
                    bullet_points.append(txt[1:].strip())
                else:
                    content_str += txt + " "

        sections_data.append({
            "section_num": section_num,
            "section_title": section_title,
            "content": content_str.strip(),
            "bullets": bullet_points,
            "chart_lines": chart_lines,
            "table_lines": table_lines
        })

    # Sort by section number
    sections_data.sort(key=lambda x: x["section_num"])

    # ----------------------------------------------------------------
    # 5) Build a Word document with the doc_title from Gemini
    # ----------------------------------------------------------------
    doc = Document()

    # Add the doc_title as a heading
    heading_para = doc.add_heading(doc_title, 0)
    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for sdata in sections_data:
        # Section heading
        sec_heading = doc.add_heading(sdata["section_title"], level=1)
        for run in sec_heading.runs:
            run.font.size = Pt(24)
            run.font.bold = True

        # Content
        if sdata["content"]:
            p = doc.add_paragraph(sdata["content"])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Bullet points
        for bullet in sdata["bullets"]:
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run(bullet)

        # If we have table lines, parse them into a Word table
        if sdata["table_lines"]:
            table_data = []
            for line in sdata["table_lines"]:
                row_line = line.strip("|")
                cells = [c.strip() for c in row_line.split("|")]
                table_data.append(cells)

            if len(table_data) >= 2:
                num_cols = len(table_data[0])
                table = doc.add_table(rows=1, cols=num_cols)
                table.style = "Light List Accent 1"

                # Fill headers
                hdr_cells = table.rows[0].cells
                for i, hdr_val in enumerate(table_data[0]):
                    hdr_cells[i].text = hdr_val

                # Fill data rows (skipping second row if it's dashes)
                for row_cells in table_data[2:]:
                    row = table.add_row().cells
                    for j, val in enumerate(row_cells):
                        row[j].text = val

        # Chart lines
        for cinfo in sdata["chart_lines"]:
            doc.add_paragraph(cinfo)
            match_png = re.search(r"\(([^)]+\.png)\)", cinfo)
            if match_png:
                chart_filename = match_png.group(1)
                if chart_filename in chart_file_list:
                    if os.path.exists(chart_filename):
                        doc.add_picture(chart_filename, width=Inches(5))
                    else:
                        doc.add_paragraph(f"(Chart file not found: {chart_filename})")
                else:
                    doc.add_paragraph(f"(Unrecognized chart: {chart_filename})")

    output_filename = "output_document.docx"
    doc.save(output_filename)
    print(f"Word document saved as {output_filename}.")
